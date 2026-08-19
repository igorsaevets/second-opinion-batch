#!/usr/bin/env python3
"""Census a tree of case files: what is READABLE, what is a SCAN, how many tokens.

WHY THIS EXISTS
R04-R07 all reviewed the petition WITHOUT the exhibits it cites. The model was told
"Exhibit V-25 shows X" and had no way to check. Adding the exhibits is the point of a
large-context run — but 421 PDFs across 1.1 GB is not automatically 421 PDFs of text.
A scanned passport has zero extractable characters and costs nothing to include, while a
born-digital market report can carry 200K.

THE TRAP THIS GUARDS AGAINST
If scanned exhibits are dropped SILENTLY, the reviewing model sees a record where those
exhibits do not exist and reports "Exhibit B-01 is missing" as a finding. That is a false
positive manufactured by the pipeline, not a defect in the filing. So the census emits an
explicit UNREADABLE list, which the corpus builder must ship to the model as
"these exist and were not machine-readable" rather than omitting them.

Read-only. Writes one JSON. No network, no spend.
"""
from __future__ import annotations

import argparse
import json
import pathlib
import sys
import time

try:
    import fitz  # PyMuPDF
except ImportError:                                            # pragma: no cover
    print("PyMuPDF (fitz) required", file=sys.stderr)
    raise

# chars-per-page below which a PDF is treated as a scan. A born-digital page of prose runs
# 1500-3500 chars; a scan with a stray header stamp runs 0-40. 100 sits in the empty gap,
# but the census records raw chars so the threshold can be re-judged without re-running.
SCAN_CHARS_PER_PAGE = 100

TEXT_EXT = {".md", ".txt"}
DOC_EXT = {".docx"}


def census_pdf(p: pathlib.Path) -> dict:
    rec: dict = {"path": str(p), "bytes": p.stat().st_size, "ext": ".pdf"}
    try:
        with fitz.open(p) as doc:
            rec["pages"] = doc.page_count
            chars = 0
            # cap the pages we read on monsters: a 3000-page exhibit would dominate the
            # census wall-clock and we only need the per-page density to classify it.
            probe = min(doc.page_count, 400)
            for i in range(probe):
                chars += len(doc.load_page(i).get_text("text"))
            rec["chars_probed"] = chars
            rec["pages_probed"] = probe
            rec["chars"] = int(chars * doc.page_count / probe) if probe else 0
            per_page = chars / probe if probe else 0
            rec["chars_per_page"] = round(per_page, 1)
            rec["kind"] = "TEXT" if per_page >= SCAN_CHARS_PER_PAGE else "SCAN"
    except Exception as e:                                     # noqa: BLE001
        rec["kind"] = "ERROR"
        rec["error"] = f"{type(e).__name__}: {str(e)[:200]}"
        rec["chars"] = 0
    sidecar = p.with_suffix(".md")
    rec["ocr_sidecar"] = str(sidecar) if sidecar.exists() else None
    rec["ocr_sidecar_chars"] = sidecar.stat().st_size if sidecar.exists() else 0
    return rec


def census_plain(p: pathlib.Path) -> dict:
    try:
        chars = len(p.read_text(encoding="utf-8", errors="replace"))
    except Exception as e:                                     # noqa: BLE001
        return {"path": str(p), "ext": p.suffix.lower(), "kind": "ERROR",
                "error": str(e)[:200], "chars": 0, "bytes": p.stat().st_size}
    return {"path": str(p), "ext": p.suffix.lower(), "kind": "TEXT",
            "chars": chars, "bytes": p.stat().st_size}


def census_docx(p: pathlib.Path) -> dict:
    rec = {"path": str(p), "ext": ".docx", "bytes": p.stat().st_size}
    try:
        import docx                                            # python-docx
        d = docx.Document(str(p))
        chars = sum(len(x.text) for x in d.paragraphs)
        for t in d.tables:
            for row in t.rows:
                chars += sum(len(c.text) for c in row.cells)
        rec.update(kind="TEXT", chars=chars)
    except Exception as e:                                     # noqa: BLE001
        rec.update(kind="ERROR", error=f"{type(e).__name__}: {str(e)[:160]}", chars=0)
    return rec


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--root", action="append", required=True,
                    help="repeatable; each is a directory or a single file")
    ap.add_argument("--out", required=True)
    ap.add_argument("--skip-ext", default=".pyc,.log,.ndjson,.png,.jpg,.jpeg,.zip,.exe")
    a = ap.parse_args()

    skip = {e.strip().lower() for e in a.skip_ext.split(",") if e.strip()}
    out = pathlib.Path(a.out)
    out.parent.mkdir(parents=True, exist_ok=True)

    files: list[pathlib.Path] = []
    for r in a.root:
        p = pathlib.Path(r)
        if p.is_file():
            files.append(p)
        elif p.is_dir():
            files.extend(x for x in p.rglob("*") if x.is_file())
        else:
            print(f"MISSING ROOT: {r}", file=sys.stderr)

    recs: list[dict] = []
    t0 = time.monotonic()
    for i, p in enumerate(files, 1):
        ext = p.suffix.lower()
        if ext in skip:
            continue
        if ext == ".pdf":
            recs.append(census_pdf(p))
        elif ext in TEXT_EXT or ext == ".json":
            recs.append(census_plain(p))
        elif ext in DOC_EXT:
            recs.append(census_docx(p))
        else:
            recs.append({"path": str(p), "ext": ext, "kind": "SKIP",
                         "chars": 0, "bytes": p.stat().st_size})
        if i % 50 == 0:
            print(f"  ...{i}/{len(files)}  {round(time.monotonic()-t0)}s", flush=True)

    by_kind: dict[str, dict] = {}
    for r in recs:
        k = by_kind.setdefault(r["kind"], {"files": 0, "chars": 0, "bytes": 0})
        k["files"] += 1
        k["chars"] += r.get("chars", 0)
        k["bytes"] += r.get("bytes", 0)

    scans_with_ocr = sum(1 for r in recs
                         if r["kind"] == "SCAN" and r.get("ocr_sidecar"))
    scans_without = sum(1 for r in recs
                        if r["kind"] == "SCAN" and not r.get("ocr_sidecar"))

    summary = {
        "_captured": time.strftime("%Y-%m-%d %H:%M"),
        "roots": a.root,
        "files_seen": len(files),
        "files_censused": len(recs),
        "wall_s": round(time.monotonic() - t0, 1),
        "by_kind": by_kind,
        "scans_with_ocr_sidecar": scans_with_ocr,
        "scans_without_ocr_sidecar": scans_without,
        "_chars_to_tokens_note": "divide chars by ~3.5 for a rough token count on mixed "
                                 "EN/RU legal text; MEASURE with countTokens before spending",
    }
    out.write_text(json.dumps({"summary": summary, "files": recs},
                              ensure_ascii=False, indent=1), encoding="utf-8")

    print("\n=== CENSUS ===")
    for k, v in sorted(by_kind.items()):
        print(f"  {k:6}  {v['files']:5} files  {v['chars']:12,} chars  "
              f"{v['bytes']/1e6:9.1f} MB")
    print(f"  SCAN with OCR sidecar: {scans_with_ocr}   without: {scans_without}")
    print(f"  wall {summary['wall_s']}s -> {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
