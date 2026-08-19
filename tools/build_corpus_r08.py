#!/usr/bin/env python3
"""Assemble a tiered case corpus for a large-context review run.

WHAT IS NEW HERE versus the R04-R07 builder
1. **AcroForm widget values.** `page.get_text()` returns the BLANK TEMPLATE of a PDF form,
   not what was typed into it. Measured 2026-08-19: `i-140 ... .pdf` yields 10,783 chars of
   template text and **92 filled widgets that the text layer does not contain** — the
   applicant's names, dates, tax number, addresses. A corpus built without widgets shows the
   reviewing model an EMPTY I-140 and invites "Part 1 left blank" as a fatal finding. That is
   a false positive manufactured by the pipeline. Widgets are extracted separately and
   labelled.
2. **An explicit NOT-INCLUDED manifest.** Whatever is excluded — unreadable scan, hard-excluded
   payment form, tier left out for budget — is named to the model as existing-but-unread.
   Silence reads as absence, and absence reads as a defect in the filing.
3. **Tiers.** The 1,048,576-token ceiling is hard and this material measures ~1.9-2.6M tokens,
   so the corpus is composed, not dumped. Each tier is separable so the ceiling can be met by
   dropping the cheapest-value tier rather than by truncating mid-document.

HARD EXCLUSIONS (never sent anywhere, at any tier)
G-1450 is the USCIS credit-card authorisation form: cardholder name, card number, expiry and
CVV. It carries no review value whatsoever — it authorises a fee, it makes no representation
about eligibility — and payment-card data has no business leaving the machine. Excluded by
filename, and the exclusion is reported rather than silent.

Writes into private/ (gitignored). No network, no spend.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import pathlib
import re
import sys

try:
    import fitz
except ImportError:                                            # pragma: no cover
    print("PyMuPDF (fitz) required", file=sys.stderr)
    raise

# Filename patterns that must never enter a payload. Matched case-insensitively against the
# file NAME, not the path, so a folder rename cannot silently disarm the guard.
HARD_EXCLUDE = [
    (re.compile(r"g-?1450", re.I), "USCIS credit-card authorisation: cardholder data, no review value"),
    (re.compile(r"\bcvv\b", re.I), "filename advertises card verification value"),
]

# Widget names that are machine-encoded duplicates of fields already captured in plain text.
NOISE_WIDGET = re.compile(r"barcode|PDF417", re.I)

SKIP_EXT = {".pyc", ".log", ".ndjson", ".png", ".jpg", ".jpeg", ".zip", ".exe",
            ".py", ".csv"}
# Sidecars from the OCR pipeline: bbox/confidence dumps. Enormous, and the useful part of
# them (the chosen reading) is already in the .md the pipeline emits next to the PDF.
SKIP_NAME = re.compile(r"\.(evidence|unconfirmed)\.json$", re.I)


def sha8(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()[:8]


def read_pdf(p: pathlib.Path) -> tuple[str, dict]:
    """Return (text, meta). Text carries page markers and a FORM FIELDS block if any."""
    meta: dict = {"pages": 0, "widgets": 0, "text_chars": 0}
    body: list[str] = []
    fields: list[str] = []
    with fitz.open(p) as doc:
        meta["pages"] = doc.page_count
        for i in range(doc.page_count):
            page = doc.load_page(i)
            t = page.get_text("text").strip()
            if t:
                body.append(f"<!-- page {i + 1} -->\n{t}")
            for w in (page.widgets() or []):
                name = w.field_name or f"page{i+1}_unnamed"
                if NOISE_WIDGET.search(name):
                    continue
                v = w.field_value
                if v in (None, "", False, "Off"):
                    continue
                v = "[X]" if v is True else str(v)
                fields.append(f"p{i+1} | {name} = {v}")
    meta["widgets"] = len(fields)
    text = "\n\n".join(body)
    meta["text_chars"] = len(text)
    if fields:
        # The header says loudly where these came from: a reviewer (human or model) must not
        # confuse a value typed into a form with a sentence written in a brief.
        text += ("\n\n<!-- FILLED FORM FIELDS (AcroForm widget values; these are the data "
                 "actually entered on the form and do NOT appear in the page text above) -->\n"
                 + "\n".join(fields))
    return text, meta


def read_docx(p: pathlib.Path) -> tuple[str, dict]:
    import docx
    d = docx.Document(str(p))
    parts = [x.text for x in d.paragraphs if x.text.strip()]
    for ti, t in enumerate(d.tables, 1):
        parts.append(f"<!-- table {ti} -->")
        for row in t.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts), {"paragraphs": len(d.paragraphs), "tables": len(d.tables)}


def read_any(p: pathlib.Path) -> tuple[str, dict]:
    ext = p.suffix.lower()
    if ext == ".pdf":
        return read_pdf(p)
    if ext == ".docx":
        return read_docx(p)
    return p.read_text(encoding="utf-8", errors="replace"), {}


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--tier", action="append", required=True, metavar="NAME=PATH",
                    help="repeatable; PATH is a file or a directory")
    ap.add_argument("--outdir", required=True)
    ap.add_argument("--min-chars", type=int, default=40,
                    help="below this a file is reported as UNREADABLE, not included")
    a = ap.parse_args()

    outdir = pathlib.Path(a.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    manifest: dict = {"tiers": {}, "excluded": [], "unreadable": [], "errors": []}

    for spec in a.tier:
        name, _, path = spec.partition("=")
        root = pathlib.Path(path)
        if not root.exists():
            manifest["errors"].append({"tier": name, "path": path, "error": "MISSING ROOT"})
            print(f"MISSING ROOT for tier {name}: {path}")
            continue

        files = [root] if root.is_file() else sorted(
            (x for x in root.rglob("*") if x.is_file()), key=lambda x: str(x).lower())

        chunks: list[str] = []
        recs: list[dict] = []
        for p in files:
            if p.suffix.lower() in SKIP_EXT or SKIP_NAME.search(p.name):
                continue
            hit = next(((rx, why) for rx, why in HARD_EXCLUDE if rx.search(p.name)), None)
            if hit:
                manifest["excluded"].append(
                    {"tier": name, "file": p.name, "reason": hit[1]})
                continue
            try:
                text, meta = read_any(p)
            except Exception as e:                             # noqa: BLE001
                manifest["errors"].append(
                    {"tier": name, "file": str(p), "error": f"{type(e).__name__}: {e}"[:200]})
                continue
            if len(text.strip()) < a.min_chars:
                manifest["unreadable"].append(
                    {"tier": name, "file": p.name,
                     "rel": str(p.relative_to(root)) if root.is_dir() else p.name,
                     "pages": meta.get("pages"), "chars": len(text.strip()),
                     "why": "no extractable text layer (scan) and no OCR sidecar"})
                continue
            rel = str(p.relative_to(root)) if root.is_dir() else p.name
            digest = sha8(p.read_bytes())
            chunks.append(
                f"\n<<<<<< FILE tier={name} path={rel} sha={digest} >>>>>>\n{text}\n"
                f"<<<<<< END {rel} >>>>>>\n")
            recs.append({"rel": rel, "sha8": digest, "chars": len(text), **meta})

        blob = "".join(chunks)
        dst = outdir / f"tier-{name}.txt"
        dst.write_text(blob, encoding="utf-8")
        manifest["tiers"][name] = {
            "root": str(root), "files": len(recs), "chars": len(blob),
            "out": str(dst), "widgets_total": sum(r.get("widgets", 0) for r in recs),
            "records": recs,
        }
        print(f"  tier {name:10} {len(recs):4} files  {len(blob):10,} chars  "
              f"{sum(r.get('widgets',0) for r in recs):5} form widgets  -> {dst.name}")

    (outdir / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=1), encoding="utf-8")

    tot = sum(t["chars"] for t in manifest["tiers"].values())
    print(f"\n  TOTAL {tot:,} chars across {len(manifest['tiers'])} tiers")
    print(f"  hard-excluded : {len(manifest['excluded'])}  "
          f"(these are NAMED to the model, not hidden)")
    print(f"  unreadable    : {len(manifest['unreadable'])}")
    print(f"  errors        : {len(manifest['errors'])}")
    for e in manifest["errors"][:10]:
        print(f"     {e.get('file', e.get('path'))}: {e['error'][:90]}")
    print(f"  -> {outdir / 'manifest.json'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
